from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_run_font(run, font_name: str, size_pt: int | float, *, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def add_title(doc: Document, title: str, question_count: int) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(6)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, "Calibri", 16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run(f"Total soal: {question_count}")
    set_run_font(subtitle_run, "Calibri", 11)


def add_question_block(doc: Document, row: dict) -> None:
    question_paragraph = doc.add_paragraph()
    question_paragraph.paragraph_format.space_after = Pt(6)
    number_run = question_paragraph.add_run(f"{row['question_order']}. ")
    set_run_font(number_run, "Calibri", 11, bold=True)
    stem_run = question_paragraph.add_run(row["stem"])
    set_run_font(stem_run, "Calibri", 11)

    options = sorted(row["options"], key=lambda item: item.get("sort_order", 0))
    for option in options:
        option_paragraph = doc.add_paragraph()
        option_paragraph.paragraph_format.left_indent = Inches(0.2)
        option_paragraph.paragraph_format.space_after = Pt(3)
        option_run = option_paragraph.add_run(f"{option['key']}. {option['text']}")
        set_run_font(option_run, "Calibri", 11)

    answer_paragraph = doc.add_paragraph()
    answer_paragraph.paragraph_format.space_before = Pt(4)
    answer_paragraph.paragraph_format.space_after = Pt(4)
    answer_label = answer_paragraph.add_run("Jawaban: ")
    set_run_font(answer_label, "Calibri", 11, bold=True)
    answer_value = answer_paragraph.add_run((row.get("correct_option_key") or "").upper())
    set_run_font(answer_value, "Calibri", 11)

    explanation_paragraph = doc.add_paragraph()
    explanation_paragraph.paragraph_format.space_after = Pt(12)
    explanation_label = explanation_paragraph.add_run("Pembahasan: ")
    set_run_font(explanation_label, "Calibri", 11, bold=True)
    explanation_value = explanation_paragraph.add_run(row.get("explanation_text") or "")
    set_run_font(explanation_value, "Calibri", 11)


def set_cell_text(cell, text: str, *, bold: bool = False, align_center: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "Calibri", 11, bold=bold)


def append_cell_paragraph(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "Calibri", 11, bold=bold)


def add_table_export(doc: Document, rows: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(0.5), Inches(1.3), Inches(3.9), Inches(2.3)]
    headers = ["No", "KISI-KISI", "SOAL", "PEMBAHASAN"]

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].width = widths[index]
        header_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_text(header_cells[index], header, bold=True, align_center=True)

    for row in rows:
        cells = table.add_row().cells
        for index, width in enumerate(widths):
            cells[index].width = width
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        set_cell_text(cells[0], str(row["question_order"]), align_center=True)
        set_cell_text(cells[1], "")

        soal_lines = [row["stem"]]
        for option in sorted(row["options"], key=lambda item: item.get("sort_order", 0)):
            soal_lines.append(f"{option['key']}. {option['text']}")
        set_cell_text(cells[2], soal_lines[0])
        for line in soal_lines[1:]:
            append_cell_paragraph(cells[2], line)

        set_cell_text(cells[3], f"Jawaban: {(row.get('correct_option_key') or '').upper()}")
        append_cell_paragraph(cells[3], "")
        append_cell_paragraph(cells[3], f"Pembahasan: {row.get('explanation_text') or ''}")


def main(argv: list[str]) -> int:
    if len(argv) not in (3, 4):
        print(
            "Usage: python scripts/export_scheduled_event_to_docx.py <input.json> <output.docx> [paragraph|table]",
            file=sys.stderr,
        )
        return 1

    input_path = Path(argv[1]).resolve()
    output_path = Path(argv[2]).resolve()
    layout = argv[3].strip().lower() if len(argv) == 4 else "paragraph"

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    payload = json.loads(input_path.read_text(encoding="utf-8-sig"))
    rows = payload.get("rows", [])
    if not rows:
        print("No rows found in input JSON.", file=sys.stderr)
        return 1

    event_title = rows[0].get("title") or "Scheduled Event Export"

    doc = Document()
    configure_document(doc)
    add_title(doc, event_title, len(rows))

    if layout == "table":
        add_table_export(doc, rows)
    else:
        for row in rows:
            add_question_block(doc, row)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
