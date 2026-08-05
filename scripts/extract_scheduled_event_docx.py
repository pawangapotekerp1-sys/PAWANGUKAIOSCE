from __future__ import annotations

import csv
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document


QUESTION_COUNT = 5
ANSWER_RE = re.compile(r"jawaban\s*:?\s*([A-E])\b", re.IGNORECASE)
EXPLANATION_LABEL_RE = re.compile(
    r"^(penjelasan|pembahasan|alasan|penjelasam)\s*:?\s*",
    re.IGNORECASE,
)
OPTION_PREFIX_RE = re.compile(r"^[A-Ea-e][\.\)]\s*")


def normalize_text(value: str) -> str:
    cleaned = (
        value.replace("\r", "")
        .replace("\xa0", " ")
        .replace("\u200b", "")
        .replace("\u00bd", "1/2")
        .replace("\u00b5g", "mcg")
        .replace("\u03bcg", "mcg")
        .replace("“", '"')
        .replace("”", '"')
        .replace("’", "'")
        .replace("–", "-")
        .replace("—", "-")
    )
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def split_non_empty_lines(value: str) -> list[str]:
    return [line.strip() for line in normalize_text(value).split("\n") if line.strip()]


def parse_question_cell(raw_value: str) -> tuple[str, list[str], list[str]]:
    lines = split_non_empty_lines(raw_value)
    notes: list[str] = []

    if len(lines) < QUESTION_COUNT + 1:
        raise ValueError(f"Question cell does not contain enough lines to split stem and {QUESTION_COUNT} options.")

    options = [OPTION_PREFIX_RE.sub("", line).strip() for line in lines[-QUESTION_COUNT:]]
    stem_lines = lines[:-QUESTION_COUNT]

    if len(options) != QUESTION_COUNT:
        notes.append("option_count_not_equal_to_five")

    return "\n".join(stem_lines).strip(), options, notes


def parse_explanation_cell(raw_value: str) -> tuple[str | None, str, list[str]]:
    lines = split_non_empty_lines(raw_value)
    notes: list[str] = []

    if not lines:
        return None, "", ["empty_explanation_cell"]

    answer_match = ANSWER_RE.search(lines[0])
    if not answer_match:
        notes.append("answer_key_not_found_in_first_line")
    answer_key = answer_match.group(1).upper() if answer_match else None

    explanation_lines = lines[1:]

    while explanation_lines and EXPLANATION_LABEL_RE.match(explanation_lines[0]):
        explanation_lines[0] = EXPLANATION_LABEL_RE.sub("", explanation_lines[0]).strip()
        if not explanation_lines[0]:
            explanation_lines = explanation_lines[1:]
        else:
            break

    explanation = "\n".join(line for line in explanation_lines if line).strip()
    if not explanation:
        notes.append("empty_explanation_after_cleanup")

    return answer_key, explanation, notes


def build_question_record(row_index: int, row_cells: list[str]) -> dict[str, Any]:
    if len(row_cells) != 4:
        raise ValueError(f"Row {row_index} does not have 4 cells.")

    number_raw, kisi_kisi_raw, question_raw, explanation_raw = row_cells

    stem, options, question_notes = parse_question_cell(question_raw)
    answer_key, explanation_text, explanation_notes = parse_explanation_cell(explanation_raw)

    notes = question_notes + explanation_notes
    for field_value in [number_raw, kisi_kisi_raw, question_raw, explanation_raw]:
        if "\ufffd" in field_value or "�" in field_value:
            notes.append("contains_replacement_character")
            break

    return {
        "sourceRowNumber": row_index,
        "questionNumberLabel": normalize_text(number_raw),
        "kisiKisi": normalize_text(kisi_kisi_raw),
        "questionText": stem,
        "options": [
            {"key": "A", "text": options[0]},
            {"key": "B", "text": options[1]},
            {"key": "C", "text": options[2]},
            {"key": "D", "text": options[3]},
            {"key": "E", "text": options[4]},
        ],
        "correctOptionKey": answer_key,
        "explanationText": explanation_text,
        "blockId": None,
        "topicId": None,
        "parsingNotes": notes,
    }


def build_csv_rows(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    csv_rows: list[dict[str, Any]] = []

    for record in records:
        option_map = {option["key"]: option["text"] for option in record["options"]}
        csv_rows.append(
            {
                "source_row_number": record["sourceRowNumber"],
                "question_number_label": record["questionNumberLabel"],
                "kisi_kisi": record["kisiKisi"],
                "question_text": record["questionText"],
                "option_a": option_map.get("A", ""),
                "option_b": option_map.get("B", ""),
                "option_c": option_map.get("C", ""),
                "option_d": option_map.get("D", ""),
                "option_e": option_map.get("E", ""),
                "correct_answer": record["correctOptionKey"] or "",
                "explanation": record["explanationText"],
                "block": "",
                "topic": "",
                "parsing_notes": "|".join(record["parsingNotes"]),
            }
        )

    return csv_rows


def build_event_update_template(records: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "targetEventTitle": "LATIHAN SOAL SBA",
        "operation": "replace_existing_questions",
        "notes": [
            "Copy the existing event metadata from production before invoking upsert_scheduled_tryout_event.",
            "This template intentionally keeps blockId and topicId null for all scheduled-event questions.",
        ],
        "payload": {
            "title": "LATIHAN SOAL SBA",
            "description": "__COPY_FROM_EXISTING_EVENT__",
            "editorialStatus": "__COPY_FROM_EXISTING_EVENT__",
            "accessStartAt": "__COPY_FROM_EXISTING_EVENT__",
            "accessEndAt": "__COPY_FROM_EXISTING_EVENT__",
            "questions": [
                {
                    "stem": record["questionText"],
                    "questionImagePath": None,
                    "blockId": None,
                    "topicId": None,
                    "correctOptionKey": record["correctOptionKey"],
                    "explanationText": record["explanationText"],
                    "explanationImagePath": None,
                    "options": record["options"],
                }
                for record in records
            ],
        },
    }


def write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    if not rows:
        raise ValueError("No rows available for CSV export.")

    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: python scripts/extract_scheduled_event_docx.py <input.docx> <output_dir>", file=sys.stderr)
        return 1

    input_path = Path(argv[1]).resolve()
    output_dir = Path(argv[2]).resolve()

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    document = Document(str(input_path))
    if not document.tables:
        print("The DOCX does not contain any table to parse.", file=sys.stderr)
        return 1

    table = document.tables[0]
    body_rows = table.rows[1:]
    records = [
        build_question_record(
            row_index=index,
            row_cells=[cell.text for cell in row.cells],
        )
        for index, row in enumerate(body_rows, start=1)
    ]

    output_dir.mkdir(parents=True, exist_ok=True)

    csv_rows = build_csv_rows(records)
    normalized_payload = {
        "sourceDocument": str(input_path),
        "questionCount": len(records),
        "questions": records,
    }
    template_payload = build_event_update_template(records)
    summary_payload = {
        "sourceDocument": str(input_path),
        "questionCount": len(records),
        "rowsWithParsingNotes": [
            {
                "sourceRowNumber": record["sourceRowNumber"],
                "questionNumberLabel": record["questionNumberLabel"],
                "parsingNotes": record["parsingNotes"],
            }
            for record in records
            if record["parsingNotes"]
        ],
    }

    write_csv(output_dir / "latihan-soal-sba-staging.csv", csv_rows)
    write_json(output_dir / "latihan-soal-sba-normalized.json", normalized_payload)
    write_json(output_dir / "latihan-soal-sba-scheduled-event-update.template.json", template_payload)
    write_json(output_dir / "latihan-soal-sba-parse-summary.json", summary_payload)

    print(f"Parsed {len(records)} questions from {input_path.name}")
    print(f"Output directory: {output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
